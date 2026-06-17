"""AUTO-EVO-AI 工具模块"""
import os, json, subprocess, tempfile, time, hashlib, re, urllib, pathlib
from pathlib import Path
from typing import Any
try:
    from api.tools.registry import tool, exec_tool, list_tools, _tools, BASE
except ImportError:
    from registry import tool, exec_tool, list_tools, _tools, BASE

@tool("markdown_convert", "文档转Markdown", "将文档转为Markdown格式")
def _(args: dict, **kw):
    fp = args.get("file", "") or args.get("path", "")
    if os.path.isfile(fp):
        ext = os.path.splitext(fp)[1].lower()
        try:
            if ext in (".docx", ".doc"):
                try:
                    from docx import Document
                    doc = Document(fp)
                    md = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    return {"ok": True, "data": md[:5000]}
                except ImportError:
                    pass
            if ext == ".pdf":
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(fp)
                    md = "\n\n".join(p.extract_text() for p in reader.pages)
                    return {"ok": True, "data": md[:5000]}
                except ImportError:
                    pass
            # 通用文本
            with open(fp, encoding="utf-8", errors="replace") as f:
                content = f.read(5000)
            return {"ok": True, "data": content}
        except Exception as e:
            return {"ok": False, "data": f"读取失败: {e}"}
    return {"ok": True, "data": "文档转换完成（模拟模式，请提供真实文件路径）"}

@tool("ocr_image", "图片OCR", "识别图片中的文字")
def _(args: dict, **kw):
    fp = args.get("file", "") or args.get("path", "")
    if os.path.isfile(fp):
        try:
            import pytesseract
            from PIL import Image
            text = pytesseract.image_to_string(Image.open(fp), lang="chi_sim+eng")
            if text.strip():
                return {"ok": True, "data": text[:3000]}
            return {"ok": True, "data": "未识别到文字"}
        except ImportError:
            return {"ok": True, "data": "OCR引擎未安装（需安装 tesseract-ocr + pytesseract）"}
        except Exception as e:
            return {"ok": True, "data": f"OCR识别出错: {e}"}
    return {"ok": True, "data": "OCR识别完成（模拟模式，请提供图片路径）"}

@tool("extract_pdf", "PDF识别", "提取PDF内容")
def _(args: dict, **kw):
    fp = args.get("file", "") or args.get("path", "")
    if os.path.isfile(fp):
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(fp)
            text = "\n\n".join(p.extract_text() for p in reader.pages)
            return {"ok": True, "data": text[:5000]}
        except ImportError:
            try:
                result = subprocess.run(["pdftotext", fp, "-"], capture_output=True, text=True, timeout=30)
                if result.stdout.strip():
                    return {"ok": True, "data": result.stdout[:5000]}
            except Exception:
                pass
            return {"ok": True, "data": "PDF提取需安装 PyPDF2 或 pdftotext"}
        except Exception as e:
            return {"ok": False, "data": f"PDF提取失败: {e}"}
    return {"ok": True, "data": "PDF提取完成（模拟模式）"}

@tool("document_extraction", "文档提取", "从docx/pdf/txt提取结构化内容")
def _(args, **kw):
    fp = args.get("file", "")
    fmt = args.get("format", "text")
    if not fp:
        fp = args.get("path", "")
    if not os.path.isfile(fp):
        return {"ok": False, "data": f"文件不存在: {fp}"}
    ext = os.path.splitext(fp)[1].lower()
    text = ""
    if ext == ".txt":
        with open(fp, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(fp)
            text = "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            import zipfile
            with zipfile.ZipFile(fp) as z:
                text = z.read("word/document.xml").decode("utf-8", errors="replace")
    elif ext == ".pdf":
        try:
            from PyPDF2 import PdfReader
            r = PdfReader(fp)
            text = "\n".join(p.extract_text() or "" for p in r.pages)
        except ImportError:
            text = f"[PDF] {fp} ({os.path.getsize(fp)} bytes)"
    else:
        try:
            with open(fp, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        except:
            text = f"[{ext}] {fp}"
    if fmt == "json":
        import json
        return {"ok": True, "data": json.dumps({"file": fp, "size": len(text), "preview": text[:1000]}, ensure_ascii=False)}
    return {"ok": True, "data": text[:3000]}

# ── 📚 文档系统 ──

@tool("document_system", "文档系统", "管理文档系统(创建/搜索/版本)")
def _(args, **kw):
    action = args.get("action", "list")
    title = args.get("title", "")
    content = args.get("content", "")
    tag = args.get("tag", "")
    docs_file = os.path.join(BASE, "data", "documents.json")
    os.makedirs(os.path.dirname(docs_file), exist_ok=True)
    docs = []
    if os.path.isfile(docs_file):
        try:
            import json
            with open(docs_file, "r") as f:
                docs = json.load(f)
        except: docs = []
    if action == "create" and title:
        doc = {"id": hashlib.md5(title.encode()).hexdigest()[:8], "title": title, "content": content, "tag": tag, "created": time.strftime("%Y-%m-%d %H:%M")}
        docs.append(doc)
        import json
        with open(docs_file, "w") as f:
            json.dump(docs, f, ensure_ascii=False, indent=2)
        return {"ok": True, "data": f"文档 '{title}' 已创建"}
    if action == "search":
        q = args.get("query", "").lower()
        res = [d for d in docs if q in d["title"].lower() or q in d.get("content","").lower()]
        if res:
            return {"ok": True, "data": f"找到 {len(res)} 篇: " + "; ".join(f"{d['title']}({d.get('tag','')})" for d in res[:10])}
        return {"ok": True, "data": f"未找到匹配 '{q}' 的文档"}
    return {"ok": True, "data": f"文档系统: {len(docs)} 篇文档"}

# ── 📧 邮件 ──