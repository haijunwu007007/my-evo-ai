from __future__ import annotations

"""

AUTO-EVO-AI V0.1 — Word文档处理器 (python-docx)

Grade: A (生产级) | Category: documents

职责：创建/编辑/转换 Word 文档 (.docx)，支持样式、表格、图片、页眉页脚、查找替换、合并

"""



__module_meta__ = {

    "id": "docx-processor",

    "name": "Docx Processor",

    "version": "V0.1",

    "group": "documents",

    "inputs": [

        {"name": "action", "type": "string", "required": True, "description": "create / edit / read / merge / replace"},

        {"name": "path", "type": "string", "required": False, "description": "文件路径"},

        {"name": "data", "type": "dict", "required": False, "description": "操作参数"},

    ],

    "outputs": [

        {"name": "result", "type": "dict", "description": "执行结果"},

    ],

    "triggers": [],

    "depends_on": [],

    "tags": ["adapter", "doc"],

    "grade": "A",

    "description": "Word文档全功能处理器：创建、编辑、合并、查找替换，支持样式/表格/图片/页眉页脚",

}



import os

import re

import json

import time

import uuid

import logging

import asyncio

import tempfile

from pathlib import Path

from datetime import datetime

from typing import Any



try:

    from modules._base.enterprise_module import EnterpriseModule

    from modules._base.metrics import metrics_collector

    from modules._base.audit import AuditLogger

except ImportError:

    import sys

    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

    from _base.enterprise_module import EnterpriseModule

    from _base.metrics import metrics_collector

    from _base.audit import AuditLogger



logger = logging.getLogger("evo.docx")



try:

    from docx import Document

    from docx.shared import Inches, Pt, Cm, Emu, RGBColor

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    from docx.enum.table import WD_TABLE_ALIGNMENT

    from docx.enum.section import WD_ORIENT

    from docx.oxml.ns import qn, nsdecls

    from docx.oxml import parse_xml

    HAS_DOCX = True

except ImportError:

    HAS_DOCX = False

    logger.warning("python-docx not installed; docx_processor disabled")





class DocxProcessor(EnterpriseModule):

    """Word文档全功能处理器"""



    def __init__(self, *args: Any, **kwargs: Any):

        super().__init__(*args, **kwargs)

        self._name = "docx-processor"

        self._status = "ready" if HAS_DOCX else "error"

        self._work_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "data", "documents")

        os.makedirs(self._work_dir, exist_ok=True)

        self._audit = AuditLogger()

        self.stats = {"calls": 0, "errors": 0, "docs_created": 0, "docs_edited": 0, "docs_merged": 0}



    # ── 生命周期 ─────────────────────────────────────────



    def initialize(self) -> None:

        logger.info(f"[docx] initialized, HAS_DOCX={HAS_DOCX}, work_dir={self._work_dir}")



    def health_check(self) -> dict[str, Any]:

        libs = {"python-docx": HAS_DOCX}

        return {

            "status": self._status,

            "module": "docx-processor",

            "libraries": libs,

            "work_dir_exists": os.path.isdir(self._work_dir),

            "stats": self.stats,

        }



    def shutdown(self) -> None:

        logger.info("[docx] shutdown complete")



    # ── 主入口 ───────────────────────────────────────────



    async def execute(self, action: str, path: str = "", data: dict | None = None, **kwargs: Any) -> dict[str, Any]:

        if not HAS_DOCX:

            return {"status": "error", "message": "python-docx 未安装"}

        self.stats["calls"] += 1

        data = data or {}

        action_map = {

            "create": self._create_docx,

            "edit": self._edit_docx,

            "read": self._read_docx,

            "merge": self._merge_docx,

            "replace": self._replace_text,

            "list": self._list_documents,

            "to_md": self._convert_to_markdown,

            "to_txt": self._convert_to_text,

            "info": self._get_document_info,

        }

        handler = action_map.get(action)

        if not handler:

            return {"status": "error", "message": f"未知 action: {action}，可选: {list(action_map.keys())}"}

        try:

            return await handler(path, data)

        except Exception as e:

            self.stats["errors"] += 1

            logger.exception(f"[docx] execute {action} 失败")

            return {"status": "error", "message": str(e)}



    # ── 创建文档 ─────────────────────────────────────────



    async def _create_docx(self, path: str, data: dict) -> dict[str, Any]:

        """创建 Word 文档"""

        filename = data.get("filename", f"doc_{uuid.uuid4().hex[:8]}.docx")

        save_path = path or os.path.join(self._work_dir, filename)

        title = data.get("title", "文档标题")

        author = data.get("author", "AUTO-EVO-AI")

        paragraphs = data.get("paragraphs", [])

        table_data = data.get("table")

        header_text = data.get("header")

        footer_text = data.get("footer")



        doc = Document()



        # ── 页面设置 ──

        section = doc.sections[0]

        section.top_margin = Cm(2.54)

        section.bottom_margin = Cm(2.54)

        section.left_margin = Cm(3.17)

        section.right_margin = Cm(3.17)



        # ── 元数据 ──

        core_props = doc.core_properties

        core_props.title = title

        core_props.author = author

        core_props.created = datetime.now()



        # ── 页眉 ──

        if header_text:

            header = section.header

            hp = header.paragraphs[0]

            hp.text = header_text

            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in hp.runs:

                run.font.size = Pt(9)

                run.font.color.rgb = RGBColor(128, 128, 128)



        # ── 页脚 ──

        if footer_text:

            footer = section.footer

            fp = footer.paragraphs[0]

            fp.text = footer_text.replace("{page}", "1").replace("{total}", "1")

            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER



        # ── 标题 ──

        doc.add_heading(title, level=0)



        # ── 正文段落 ──

        for p_data in paragraphs:

            if isinstance(p_data, str):

                p_data = {"text": p_data, "style": "Normal"}

            text = p_data.get("text", "")

            style = p_data.get("style", "Normal")

            bold = p_data.get("bold", False)

            italic = p_data.get("italic", False)

            size = p_data.get("size")

            color = p_data.get("color")

            align = p_data.get("align")

            heading_level = p_data.get("heading")



            if heading_level is not None:

                p = doc.add_heading(text, level=heading_level)

            else:

                p = doc.add_paragraph(text, style=style)

                if bold:

                    for run in p.runs:

                        run.bold = True

                if italic:

                    for run in p.runs:

                        run.italic = True

                if size:

                    for run in p.runs:

                        run.font.size = Pt(size)

                if color:

                    for run in p.runs:

                        run.font.color.rgb = RGBColor(*self._parse_color(color))

                if align:

                    align_map = {

                        "left": WD_ALIGN_PARAGRAPH.LEFT,

                        "center": WD_ALIGN_PARAGRAPH.CENTER,

                        "right": WD_ALIGN_PARAGRAPH.RIGHT,

                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,

                    }

                    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)



        # ── 表格 ──

        if table_data:

            headers = table_data.get("headers", [])

            rows = table_data.get("rows", [])

            if headers:

                table = doc.add_table(rows=1 + len(rows), cols=len(headers))

                table.style = "Light Grid Accent 1"

                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, h in enumerate(headers):

                    cell = table.rows[0].cells[i]

                    cell.text = str(h)

                    for paragraph in cell.paragraphs:

                        for run in paragraph.runs:

                            run.bold = True

                for ri, row_data in enumerate(rows):

                    for ci, val in enumerate(row_data):

                        table.rows[ri + 1].cells[ci].text = str(val) if val else ""

                doc.add_paragraph()



        # ── 图片 ──

        image_paths = data.get("images", [])

        for img_path in image_paths:

            if os.path.isfile(img_path):

                try:

                    doc.add_picture(img_path, width=Inches(5.5))

                    doc.last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    doc.add_paragraph()

                except Exception as e:

                    logger.warning(f"[docx] 图片插入失败 {img_path}: {e}")



        doc.save(save_path)

        self.stats["docs_created"] += 1

        self._audit.log(action="create_docx", detail=f"title={title}")

        return {"status": "success", "path": os.path.abspath(save_path), "filename": filename}



    # ── 编辑文档 ─────────────────────────────────────────



    async def _edit_docx(self, path: str, data: dict) -> dict[str, Any]:

        """编辑已有 Word 文档"""

        if not path or not os.path.isfile(path):

            return {"status": "error", "message": f"文件不存在: {path}"}



        doc = Document(path)

        edits = data.get("edits", [])



        for edit in edits:

            op = edit.get("op", "")

            if op == "add_paragraph":

                text = edit.get("text", "")

                style = edit.get("style", "Normal")

                doc.add_paragraph(text, style=style)

            elif op == "add_heading":

                text = edit.get("text", "")

                level = edit.get("level", 1)

                doc.add_heading(text, level=level)

            elif op == "add_table":

                headers = edit.get("headers", [])

                rows = edit.get("rows", [])

                if headers:

                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))

                    table.style = "Light Grid Accent 1"

                    for i, h in enumerate(headers):

                        table.rows[0].cells[i].text = str(h)

                    for ri, row_data in enumerate(rows):

                        for ci, val in enumerate(row_data):

                            table.rows[ri + 1].cells[ci].text = str(val) if val else ""

            elif op == "add_image":

                img_path = edit.get("path", "")

                if os.path.isfile(img_path):

                    width = edit.get("width", 5.5)

                    doc.add_picture(img_path, width=Inches(width))

            elif op == "set_header":

                for section in doc.sections:

                    header = section.header

                    hp = header.paragraphs[0]

                    hp.text = edit.get("text", "")

            elif op == "set_footer":

                for section in doc.sections:

                    footer = section.footer

                    fp = footer.paragraphs[0]

                    fp.text = edit.get("text", "")



        doc.save(path)

        self.stats["docs_edited"] += 1

        return {"status": "success", "path": os.path.abspath(path)}



    # ── 读取文档 ─────────────────────────────────────────



    async def _read_docx(self, path: str, data: dict) -> dict[str, Any]:

        """读取 Word 文档内容"""

        if not path or not os.path.isfile(path):

            return {"status": "error", "message": f"文件不存在: {path}"}



        doc = Document(path)

        result = {

            "paragraphs": [],

            "tables": [],

            "sections": len(doc.sections),

        }



        for p in doc.paragraphs:

            result["paragraphs"].append({

                "text": p.text,

                "style": p.style.name if p.style else None,

                "bold": any(r.bold for r in p.runs) if p.runs else False,

                "heading": p.style.name.startswith("Heading") if p.style else False,

            })



        for table in doc.tables:

            rows_data = []

            for row in table.rows:

                rows_data.append([cell.text for cell in row.cells])

            result["tables"].append(rows_data)



        return {"status": "success", "content": result}



    # ── 合并文档 ─────────────────────────────────────────



    async def _merge_docx(self, path: str, data: dict) -> dict[str, Any]:

        """合并多个 Word 文档"""

        files = data.get("files", [])

        if isinstance(files, str):

            files = [files]

        if not files:

            return {"status": "error", "message": "未指定待合并文件列表"}



        save_path = path or os.path.join(self._work_dir, f"merged_{uuid.uuid4().hex[:8]}.docx")



        merged = Document()

        for fi, fp in enumerate(files):

            if not os.path.isfile(fp):

                logger.warning(f"[docx] 合并跳过不存在的文件: {fp}")

                continue

            if fi > 0:

                merged.add_page_break()

            sub = Document(fp)

            for p in sub.paragraphs:

                new_p = merged.add_paragraph(p.text, style=p.style.name if p.style else None)

                new_p.alignment = p.alignment

            for table in sub.tables:

                headers = [cell.text for cell in table.rows[0].cells]

                if headers:

                    t = merged.add_table(rows=len(table.rows), cols=len(headers))

                    t.style = table.style

                    for ri, row in enumerate(table.rows):

                        for ci, cell in enumerate(row.cells):

                            t.rows[ri].cells[ci].text = cell.text



        merged.save(save_path)

        self.stats["docs_merged"] += 1

        return {"status": "success", "path": os.path.abspath(save_path)}



    # ── 查找替换 ─────────────────────────────────────────



    async def _replace_text(self, path: str, data: dict) -> dict[str, Any]:

        """文档内查找/替换文本"""

        if not path or not os.path.isfile(path):

            return {"status": "error", "message": f"文件不存在: {path}"}



        old = data.get("old", "")

        new = data.get("new", "")

        if not old:

            return {"status": "error", "message": "缺少 old 参数"}



        doc = Document(path)

        count = 0



        for p in doc.paragraphs:

            for run in p.runs:

                if old in run.text:

                    run.text = run.text.replace(old, new)

                    count += 1



        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for p in cell.paragraphs:

                        for run in p.runs:

                            if old in run.text:

                                run.text = run.text.replace(old, new)

                                count += 1



        doc.save(path)

        return {"status": "success", "replacements": count}



    # ── 转换为 Markdown ──────────────────────────────────



    async def _convert_to_markdown(self, path: str, data: dict) -> dict[str, Any]:

        """将 Word 文档转换为 Markdown"""

        if not path or not os.path.isfile(path):

            return {"status": "error", "message": f"文件不存在: {path}"}



        doc = Document(path)

        lines = ["# " + (doc.core_properties.title or "Untitled"), ""]



        for p in doc.paragraphs:

            style_name = p.style.name if p.style else ""

            text = p.text.strip()

            if not text:

                continue

            if style_name.startswith("Heading 1"):

                lines.append(f"# {text}")

            elif style_name.startswith("Heading 2"):

                lines.append(f"## {text}")

            elif style_name.startswith("Heading 3"):

                lines.append(f"### {text}")

            else:

                lines.append(text)

            lines.append("")



        for table in doc.tables:

            headers = [cell.text for cell in table.rows[0].cells]

            lines.append("| " + " | ".join(headers) + " |")

            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

            for row in table.rows[1:]:

                cells = [cell.text for cell in row.cells]

                lines.append("| " + " | ".join(cells) + " |")

            lines.append("")



        md = "\n".join(lines)

        if data.get("save", False):

            md_path = path.replace(".docx", ".md") if path.endswith(".docx") else path + ".md"

            with open(md_path, "w", encoding="utf-8") as f:

                f.write(md)

            return {"status": "success", "path": os.path.abspath(md_path), "content": md}



        return {"status": "success", "content": md}



    # ── 转换为纯文本 ─────────────────────────────────────



    async def _convert_to_text(self, path: str, data: dict) -> dict[str, Any]:

        """将 Word 文档转换为纯文本"""

        if not path or not os.path.isfile(path):

            return {"status": "error", "message": f"文件不存在: {path}"}



        doc = Document(path)

        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())



        if data.get("save", False):

            txt_path = path.replace(".docx", ".txt") if path.endswith(".docx") else path + ".txt"

            with open(txt_path, "w", encoding="utf-8") as f:

                f.write(text)

            return {"status": "success", "path": os.path.abspath(txt_path), "text": text}



        return {"status": "success", "text": text}



    # ── 文档信息 ─────────────────────────────────────────



    async def _get_document_info(self, path: str, data: dict) -> dict[str, Any]:

        """获取文档元信息"""

        if not path or not os.path.isfile(path):

            return {"status": "error", "message": f"文件不存在: {path}"}



        doc = Document(path)

        cp = doc.core_properties

        info = {

            "title": cp.title or "",

            "author": cp.author or "",

            "created": str(cp.created or ""),

            "modified": str(cp.modified or ""),

            "paragraphs": len(doc.paragraphs),

            "tables": len(doc.tables),

            "sections": len(doc.sections),

            "file_size_bytes": os.path.getsize(path),

        }

        return {"status": "success", "info": info}



    # ── 列出文档目录 ─────────────────────────────────────



    async def _list_documents(self, path: str, data: dict) -> dict[str, Any]:

        """列出工作目录下的文档"""

        search_dir = path or self._work_dir

        if not os.path.isdir(search_dir):

            return {"status": "error", "message": f"目录不存在: {search_dir}"}



        files = []

        for f in sorted(os.listdir(search_dir)):

            fp = os.path.join(search_dir, f)

            if os.path.isfile(fp) and f.lower().endswith(".docx"):

                files.append({

                    "name": f,

                    "size_bytes": os.path.getsize(fp),

                    "modified": datetime.fromtimestamp(os.path.getmtime(fp)).isoformat(),

                })

        return {"status": "success", "files": files, "directory": search_dir}



    # ── 辅助函数 ─────────────────────────────────────────



    @staticmethod

    def _parse_color(color: str) -> tuple:

        """解析颜色值，支持 #RRGGBB 或 (r,g,b)"""

        if isinstance(color, str) and color.startswith("#"):

            color = color.lstrip("#")

            return tuple(int(color[i : i + 2], 16) for i in (0, 2, 4))

        if isinstance(color, (list, tuple)) and len(color) == 3:

            return tuple(color)

        return (0, 0, 0)





module_class = DocxProcessor

