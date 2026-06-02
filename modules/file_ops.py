"""文件/邮件操作模块 — Excel读写/Word生成/邮件发送"""

import os, csv, io, smtplib, email.utils
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from pathlib import Path
from core.logging_config import get_logger

logger = get_logger("evo.file_ops")

# ── Excel ────────────────────────────

def excel_read(path: str) -> str:
    """读取 Excel 返回摘要"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        headers = [str(c) for c in (rows[0] if rows else [])]
        data_rows = len(rows) - 1
        cols = len(headers)
        sample = ", ".join(str(c)[:20] for c in (rows[1] if len(rows) > 1 else []))
        return f"📊 Excel: {os.path.basename(path)}\n  表头: {headers[:8]}\n  数据: {data_rows}行 × {cols}列\n  示例: {sample}"
    except Exception as e:
        return f"⚠️ 读取失败: {e}"

def excel_write(path: str, data: list[list]) -> str:
    """写入 Excel"""
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        for row in data:
            ws.append(row)
        wb.save(path)
        return f"✅ 已写入 {len(data)}行 → {path}"
    except Exception as e:
        return f"⚠️ 写入失败: {e}"

# ── Word ─────────────────────────────

def word_create(path: str, title: str, content: str) -> str:
    """生成 Word 文档"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, 0)
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.save(path)
        return f"✅ Word 已生成 → {path}"
    except Exception as e:
        return f"⚠️ 生成失败: {e}"

# ── 邮件 ─────────────────────────────

def send_email(smtp_host: str, smtp_port: int, user: str, pwd: str,
               to: str, subject: str, body: str, attach: str = "") -> str:
    """发送邮件"""
    try:
        msg = MIMEMultipart()
        msg["From"] = user
        msg["To"] = to
        msg["Subject"] = subject
        msg["Date"] = email.utils.formatdate()
        msg.attach(MIMEText(body, "plain", "utf-8"))
        if attach and os.path.isfile(attach):
            with open(attach, "rb") as f:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(attach)}")
            msg.attach(part)
        with smtplib.SMTP(smtp_host, smtp_port, timeout=30) as s:
            s.starttls()
            s.login(user, pwd)
            s.send_message(msg)
        return f"✅ 邮件已发送 → {to}"
    except Exception as e:
        return f"⚠️ 发送失败: {e}"

# ── CLI 入口 ─────────────────────────

def handle_file_cmd(args: list[str]) -> str:
    if not args:
        return "用法: file read <路径> | file write <路径> <内容> | word <路径> <标题> <内容> | mail <收件人> <主题> <内容>"
    cmd = args[0]
    if cmd == "read" and len(args) >= 2:
        p = " ".join(args[1:])
        if p.endswith((".xlsx", ".xls")):
            return excel_read(p)
        return f"不支持格式: {p}"
    if cmd == "write" and len(args) >= 3:
        p = args[1]
        d = [row.split(",") for row in " ".join(args[2:]).split(";")]
        return excel_write(p, d)
    if cmd == "word" and len(args) >= 3:
        p = args[1]
        t = args[2]
        c = " ".join(args[3:]) if len(args) > 3 else ""
        return word_create(p, t, c)
    if cmd == "mail" and len(args) >= 4:
        t = args[1]
        s = args[2]
        b = " ".join(args[3:])
        cfg = os.environ.get("EVO_SMTP_CONFIG", "")
        if not cfg:
            return "⚠️ 未配置 SMTP。设置环境变量 EVO_SMTP_CONFIG=host:port:user:pass"
        parts = cfg.split(":")
        if len(parts) >= 4:
            return send_email(parts[0], int(parts[1]), parts[2], parts[3], t, s, b)
        return "⚠️ SMTP 配置格式错误"
    return "用法见: file help"
